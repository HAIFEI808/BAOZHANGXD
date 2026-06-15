#!/usr/bin/env python3
"""将论文Markdown文件转换为格式化的Word文档（.docx），含自动分级目录"""

import re
import os
from docx import Document
from docx.shared import Pt, Cm, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# 附录图片路径
IMG_DIR = r'D:\cxdownload\neckedf401re-main\neckedf401re-main'
IMG_APPENDIX1 = os.path.join(IMG_DIR, 'appendix1_task_architecture.png')
IMG_APPENDIX2 = os.path.join(IMG_DIR, 'appendix2_gantt_chart.png')
IMG_FIG2_1 = os.path.join(IMG_DIR, 'fig2_1_system_architecture.png')

# ── 工具函数 ──────────────────────────────────────────────

def set_cell_shading(cell, color):
    shading = OxmlElement('w:shd')
    shading.set(qn('w:fill'), color)
    shading.set(qn('w:val'), 'clear')
    cell._tc.get_or_add_tcPr().append(shading)

def set_table_borders(table):
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
    borders = OxmlElement('w:tblBorders')
    for name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        b = OxmlElement(f'w:{name}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'), '4')
        b.set(qn('w:space'), '0')
        b.set(qn('w:color'), '000000')
        borders.append(b)
    tblPr.append(borders)

def add_body_para(doc, text, font_name='宋体', font_size=12, bold=False,
                  alignment=None, first_line_indent=0.74, space_before=Pt(2), space_after=Pt(4)):
    """正文段落"""
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.font.name = font_name
    run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    run.font.size = Pt(font_size)
    run.bold = bold
    if alignment is not None:
        p.alignment = alignment
    p.paragraph_format.space_before = space_before
    p.paragraph_format.space_after = space_after
    if first_line_indent:
        p.paragraph_format.first_line_indent = Cm(first_line_indent)
    return p

def add_code_block(doc, code_text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(6)
    p.paragraph_format.left_indent = Cm(1)
    run = p.add_run(code_text)
    run.font.name = 'Consolas'
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    return p

def add_ascii_art(doc, art_text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(3)
    p.paragraph_format.space_after = Pt(3)
    p.paragraph_format.left_indent = Cm(0.5)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(art_text)
    run.font.name = 'Consolas'
    run.font.size = Pt(6.5)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    return p

def add_image(doc, img_path, width_cm=14.0, caption=None):
    """在文档中插入图片，居中显示，可附带题注"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(8)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run()
    if os.path.exists(img_path):
        run.add_picture(img_path, width=Cm(width_cm))
    else:
        run.add_text(f'[图片未找到: {img_path}]')
        run.font.color.rgb = RGBColor(0xFF, 0x00, 0x00)
    if caption:
        add_caption(doc, caption)
    return p

def add_caption(doc, text):
    """图/表题注"""
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(text)
    run.font.name = '宋体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after = Pt(4)
    return p

def add_cover_page(doc):
    """在文档最前面添加结课论文封面"""
    # ── 封面整体包装在一个节中，独立页码 ──

    # 空行推标题到页面中部
    for _ in range(6):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(0)

    # 学校名称
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('XX大学')
    run.font.name = '黑体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    run.font.size = Pt(28)
    run.bold = True
    p.paragraph_format.space_after = Pt(24)

    # 课程论文字样
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('课 程 结 课 论 文')
    run.font.name = '黑体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    run.font.size = Pt(26)
    run.bold = True
    p.paragraph_format.space_after = Pt(40)

    # 论文题目
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run('基于FreeRTOS的智能手表\n嵌入式系统设计与实现')
    run.font.name = '黑体'
    run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
    run.font.size = Pt(22)
    run.bold = True
    p.paragraph_format.space_after = Pt(60)

    # 信息表格（无边框）
    info_data = [
        ('课程名称', '嵌入式系统设计'),
        ('学    院', '工学部'),
        ('专    业', '人工智能'),
        ('班    级', '4班'),
        ('学    号', '2023105400440'),
        ('学生姓名', '张方宇'),
        ('指导教师', '__________'),
        ('完成日期', '2026年6月'),
    ]

    table = doc.add_table(rows=len(info_data), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    for ri, (label, value) in enumerate(info_data):
        # 标签列
        cell_l = table.rows[ri].cells[0]
        cell_l.width = Cm(3.5)
        cell_l.paragraphs[0].clear()
        run = cell_l.paragraphs[0].add_run(label + '：')
        run.font.name = '宋体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        run.font.size = Pt(14)
        run.bold = True
        cell_l.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT

        # 值列
        cell_r = table.rows[ri].cells[1]
        cell_r.width = Cm(6.0)
        cell_r.paragraphs[0].clear()
        run = cell_r.paragraphs[0].add_run(value)
        run.font.name = '宋体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        run.font.size = Pt(14)
        if '___' in value:
            run.font.color.rgb = RGBColor(0xBB, 0xBB, 0xBB)
        cell_r.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.LEFT

    # 去除封面表格边框
    tbl = table._tbl
    tblPr = tbl.tblPr if tbl.tblPr is not None else OxmlElement('w:tblPr')
    borders = OxmlElement('w:tblBorders')
    for name in ['top', 'left', 'bottom', 'right', 'insideH', 'insideV']:
        b = OxmlElement(f'w:{name}')
        b.set(qn('w:val'), 'none')
        b.set(qn('w:sz'), '0')
        b.set(qn('w:space'), '0')
        b.set(qn('w:color'), 'auto')
        borders.append(b)
    tblPr.append(borders)

    # 分页
    doc.add_page_break()

def insert_toc(doc):
    """在当前位置插入 Word 自动目录域"""
    paragraph = doc.add_paragraph()
    run = paragraph.add_run()

    fldChar_begin = OxmlElement('w:fldChar')
    fldChar_begin.set(qn('w:fldCharType'), 'begin')
    run._r.append(fldChar_begin)

    instrText = OxmlElement('w:instrText')
    instrText.set(qn('xml:space'), 'preserve')
    # TOC: 显示级别1-3, 含页码右对齐, \h 超链接, \u 使用段落大纲级别
    instrText.text = ' TOC \\o "1-3" \\h \\z \\u '
    run._r.append(instrText)

    fldChar_separate = OxmlElement('w:fldChar')
    fldChar_separate.set(qn('w:fldCharType'), 'separate')
    run._r.append(fldChar_separate)

    # 目录占位文字
    run2 = paragraph.add_run('（请在Word中右键此处 → 更新域，生成目录）')
    run2.font.name = '宋体'
    run2._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    run2.font.size = Pt(10)
    run2.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    run2.italic = True

    fldChar_end = OxmlElement('w:fldChar')
    fldChar_end.set(qn('w:fldCharType'), 'end')
    run._r.append(fldChar_end)

    # 目录后分页
    doc.add_paragraph().add_run().add_break (WD_BREAK.PAGE)

# ── 正文处理函数 ──────────────────────────────────────────

def add_heading_with_style(doc, text, level):
    """使用 Word 内置 Heading 样式添加标题（纳入目录）"""
    heading = doc.add_heading(text, level=level)
    # 设置中文字体
    for run in heading.runs:
        if level == 1:
            run.font.name = '黑体'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
            run.font.size = Pt(16)
            run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
        elif level == 2:
            run.font.name = '黑体'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
            run.font.size = Pt(14)
            run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
        elif level == 3:
            run.font.name = '黑体'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
            run.font.size = Pt(12)
            run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
    return heading

# ── 主转换逻辑 ────────────────────────────────────────────

def parse_and_convert(md_path, docx_path):
    doc = Document()

    # ── 页面设置 ──
    section = doc.sections[0]
    section.page_width = Cm(21.0)
    section.page_height = Cm(29.7)
    section.top_margin = Cm(2.54)
    section.bottom_margin = Cm(2.54)
    section.left_margin = Cm(3.18)
    section.right_margin = Cm(3.18)

    # ── 定义 Word 内置标题样式 ──
    for lvl in [1, 2, 3]:
        style_name = f'Heading {lvl}'
        style = doc.styles[style_name]
        style.font.name = '黑体'
        style.element.rPr.rFonts.set(qn('w:eastAsia'), '黑体')
        style.font.bold = True
        style.font.color.rgb = RGBColor(0, 0, 0)
        if lvl == 1:
            style.font.size = Pt(16)
            style.paragraph_format.space_before = Pt(18)
            style.paragraph_format.space_after = Pt(8)
        elif lvl == 2:
            style.font.size = Pt(14)
            style.paragraph_format.space_before = Pt(12)
            style.paragraph_format.space_after = Pt(6)
        elif lvl == 3:
            style.font.size = Pt(12)
            style.paragraph_format.space_before = Pt(8)
            style.paragraph_format.space_after = Pt(4)

    # Normal 样式
    style_normal = doc.styles['Normal']
    style_normal.font.name = '宋体'
    style_normal.font.size = Pt(12)
    style_normal.element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    style_normal.paragraph_format.first_line_indent = Cm(0.74)

    # ── 封面（第一页） ──
    add_cover_page(doc)

    # ── 读取 Markdown ──
    with open(md_path, 'r', encoding='utf-8') as f:
        lines = f.readlines()

    i = 0
    in_code_block = False
    code_buffer = []
    in_ascii_art = False
    ascii_buffer = []
    in_table = False
    table_rows = []
    table_just_finished = False
    title_added = False
    toc_inserted = False
    first_heading1_seen = False
    # 附录图片替换追踪
    in_appendix1 = False
    in_appendix2 = False
    appendix1_img_done = False
    appendix2_img_done = False
    fig2_1_img_done = False
    pending_fig2_1 = False

    while i < len(lines):
        line = lines[i].rstrip()

        # ── 代码块（附录/图中的替换为图片） ──
        if line.startswith('```'):
            if in_code_block:
                in_code_block = False
                code_str = '\n'.join(code_buffer)
                code_buffer = []
                # 图2-1：系统总体架构图 → 图片
                if pending_fig2_1 and not fig2_1_img_done:
                    add_image(doc, IMG_FIG2_1, width_cm=14.0,
                              caption='图2-1  系统总体架构图')
                    fig2_1_img_done = True
                    pending_fig2_1 = False
                # 附录中的框图 → 图片
                elif in_appendix1 and not appendix1_img_done:
                    add_image(doc, IMG_APPENDIX1, width_cm=14.5,
                              caption='附图1  FreeRTOS智能手表任务架构与通信框图')
                    appendix1_img_done = True
                elif in_appendix2 and not appendix2_img_done:
                    add_image(doc, IMG_APPENDIX2, width_cm=14.0,
                              caption='附图2  项目开发规划甘特图')
                    appendix2_img_done = True
                else:
                    add_code_block(doc, code_str)
            else:
                in_code_block = True
            i += 1
            continue
        if in_code_block:
            code_buffer.append(line)
            i += 1
            continue

        # ── ASCII 框图（附录中替换为图片） ──
        is_ascii = bool(re.match(r'^[\s]*[│├└┌┤┴┬┼─┃┏┓┗┛]+', line))
        is_diagram = bool(re.match(r'^\s{4,}[┌└│├┤┬┴┼─▌▐█▀▄▊▎►◄▼▲]', line))
        if is_ascii or is_diagram:
            if (in_appendix1 and not appendix1_img_done) or \
               (in_appendix2 and not appendix2_img_done):
                # 不收集 ASCII，直接跳过 — 紧接着会插入图片
                i += 1
                continue
            if not in_ascii_art:
                in_ascii_art = True
                ascii_buffer = []
            ascii_buffer.append(line)
            i += 1
            continue
        else:
            if in_ascii_art:
                in_ascii_art = False
                art_str = '\n'.join(ascii_buffer)
                # 附录中的 ASCII 用图片替代
                if in_appendix1 and not appendix1_img_done:
                    add_image(doc, IMG_APPENDIX1, width_cm=14.5,
                              caption='附图1  FreeRTOS智能手表任务架构与通信框图')
                    appendix1_img_done = True
                elif in_appendix2 and not appendix2_img_done:
                    add_image(doc, IMG_APPENDIX2, width_cm=14.0,
                              caption='附图2  项目开发规划甘特图')
                    appendix2_img_done = True
                else:
                    add_ascii_art(doc, art_str)
                ascii_buffer = []

        # ── 表格（支持 > | 引用前缀） ──
        is_table_line = (line.startswith('|') or line.startswith('> |') or
                         line.startswith('> |'))
        # 去掉引用前缀以便统一处理
        cleaned = line[2:] if line.startswith('> ') else line
        if is_table_line and cleaned.endswith('|'):
            if not in_table:
                in_table = True
                table_rows = []
            if re.match(r'^[\|\s\-:]+$', cleaned):
                i += 1
                continue
            cells = [c.strip() for c in cleaned.split('|')[1:-1]]
            table_rows.append(cells)
            i += 1
            continue
        else:
            if in_table:
                in_table = False
                table_just_finished = True

        # 空行
        if line.strip() == '':
            i += 1
            continue

        # ── 渲染表格 ──
        if table_just_finished and table_rows:
            table_just_finished = False
            num_cols = max(len(r) for r in table_rows)
            table = doc.add_table(rows=len(table_rows), cols=num_cols)
            table.alignment = WD_TABLE_ALIGNMENT.CENTER
            set_table_borders(table)

            for ri, rd in enumerate(table_rows):
                for ci, ct in enumerate(rd):
                    if ci < num_cols:
                        cell = table.rows[ri].cells[ci]
                        cell.paragraphs[0].clear()
                        run = cell.paragraphs[0].add_run(ct)
                        run.font.name = '宋体'
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
                        run.font.size = Pt(9)
                        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
                        if ri == 0:
                            set_cell_shading(cell, '2F5496')
                            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                            run.bold = True
                        elif ri % 2 == 0:
                            set_cell_shading(cell, 'D6E4F0')
            doc.add_paragraph()
            table_rows = []
            i += 1
            continue

        # ── 标题 # → 论文大标题（不入目录） ──
        if line.startswith('# ') and not line.startswith('## '):
            text = line[2:].strip()
            if not title_added:
                add_body_para(doc, text, font_name='黑体', font_size=22, bold=True,
                              alignment=WD_ALIGN_PARAGRAPH.CENTER, first_line_indent=None,
                              space_before=Pt(36), space_after=Pt(24))
                title_added = True
                # 紧接插入目录
                if not toc_inserted:
                    insert_toc(doc)
                    toc_inserted = True
            i += 1
            continue

        # ── 标题 ## → Heading 1（章） ──
        if line.startswith('## ') and not line.startswith('### '):
            text = line[3:].strip()
            # 跳过"目录"标题（已自动生成）和"中文摘要/英文摘要"（走特殊处理）
            if text in ['目录']:
                i += 1
                continue
            if not first_heading1_seen:
                first_heading1_seen = True
            add_heading_with_style(doc, text, level=1)
            # 跟踪附录区域
            if '附录1' in text:
                in_appendix1, in_appendix2 = True, False
            elif '附录2' in text:
                in_appendix1, in_appendix2 = False, True
            else:
                in_appendix1, in_appendix2 = False, False
            i += 1
            continue

        # ── 标题 ### → Heading 2（节） ──
        if line.startswith('### ') and not line.startswith('#### '):
            text = line[4:].strip()
            add_heading_with_style(doc, text, level=2)
            # 跟踪附录子区域
            if '附录1' in text:
                in_appendix1, in_appendix2 = True, False
            elif '附录2' in text:
                in_appendix1, in_appendix2 = False, True
            i += 1
            continue

        # ── 标题 #### → Heading 3（小节） ──
        if line.startswith('#### '):
            text = line[5:].strip()
            add_heading_with_style(doc, text, level=3)
            i += 1
            continue

        # ── 水平线 ──
        if line.strip() == '---':
            p = doc.add_paragraph()
            p.paragraph_format.space_before = Pt(6)
            p.paragraph_format.space_after = Pt(6)
            pPr = p._p.get_or_add_pPr()
            pBdr = OxmlElement('w:pBdr')
            bottom = OxmlElement('w:bottom')
            bottom.set(qn('w:val'), 'single')
            bottom.set(qn('w:sz'), '6')
            bottom.set(qn('w:space'), '1')
            bottom.set(qn('w:color'), '999999')
            pBdr.append(bottom)
            pPr.append(pBdr)
            i += 1
            continue

        # ── 图/表题注 (以 > 开头) ──
        if line.startswith('> '):
            text = line[2:].strip().replace('**', '')
            # 检测图2-1题注，下一个代码块将替换为图片
            if not fig2_1_img_done and '图2-1' in text and '系统总体架构' in text:
                pending_fig2_1 = True
                add_caption(doc, text)
            elif '图2-1' in text and fig2_1_img_done:
                # 已经插入过图片，跳过重复的题注
                pass
            else:
                add_caption(doc, text)
            i += 1
            continue

        # ── 普通段落 ──
        text = line.strip()
        text = re.sub(r'\*\*(.+?)\*\*', r'\1', text)
        text = re.sub(r'`([^`]+)`', r'\1', text)

        # 跳过 Markdown 目录项
        if text.startswith('- [') and text.endswith(')'):
            i += 1
            continue

        is_body = not any([
            text.startswith('|'), text.startswith('```'),
            text.startswith('#'), text.startswith('> '),
            text.strip() in ['---', '']
        ])

        if is_body and len(text) > 10:
            add_body_para(doc, text)
        elif is_body:
            add_body_para(doc, text, first_line_indent=None)

        i += 1

    # ── 处理末尾未刷新的 ASCII 缓冲区（附录图片） ──
    if in_ascii_art and ascii_buffer:
        if in_appendix1 and not appendix1_img_done:
            add_image(doc, IMG_APPENDIX1, width_cm=14.5,
                      caption='附图1  FreeRTOS智能手表任务架构与通信框图')
            appendix1_img_done = True
        elif in_appendix2 and not appendix2_img_done:
            add_image(doc, IMG_APPENDIX2, width_cm=14.0,
                      caption='附图2  项目开发规划甘特图')
            appendix2_img_done = True
        else:
            add_ascii_art(doc, '\n'.join(ascii_buffer))

    # ── 保存 ──
    doc.save(docx_path)
    print(f"Word 文档已生成（含自动目录域）: {docx_path}")

if __name__ == '__main__':
    md_path = r'D:\cxdownload\neckedf401re-main\neckedf401re-main\论文_基于FreeRTOS的智能手表嵌入式系统设计与实现.md'
    docx_path = r'D:\cxdownload\neckedf401re-main\neckedf401re-main\论文_基于FreeRTOS的智能手表嵌入式系统设计与实现_v3.docx'
    parse_and_convert(md_path, docx_path)
